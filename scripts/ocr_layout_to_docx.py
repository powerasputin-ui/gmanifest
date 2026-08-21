"""Reconstructs an approximate visual layout from PaddleOCR's per-line
bounding boxes into a real, editable .docx — the OCR analogue of
pdf_to_docx.py's pdf2docx-based layout reconstruction for text-layer PDFs.

Pixel coordinates from OCR are used only for *relative* clustering (deciding
which lines sit in the same row, and which of those are separate columns) —
never written as absolute positions into the document. Output is always real
paragraphs/tables, matching the "live editable text, not an image or a
floating text frame" approach already used for text-layer PDFs.

The per-page routing/clustering lives in _ocr_common.py
(ocr_layout_docx_pages()) so this script and the persistent worker
(ocr_worker.py) can't drift apart — this file just wires up a fresh set of
models for a single run and writes the .docx.

Usage: python ocr_layout_to_docx.py <input.pdf|input.png|...> <output.docx> [--pages 0,2,5]
Prints OK on success. On failure, prints the traceback to stderr and exits non-zero.
"""
import sys

from _ocr_common import LazyModel, make_layout_detector, make_ocr, make_table_pipeline, ocr_layout_docx_pages

# A table with more columns than this doesn't fit a standard portrait page
# at a readable size — cells get squeezed so narrow that content wraps onto
# many lines, which reads as "no columns, just text running downward" even
# though it's technically a table. Real-world dense manifests/spec tables
# regularly have 12-20 columns, so this is the common case, not an edge
# case, for the table-recognition path.
WIDE_TABLE_COLUMN_THRESHOLD = 6
WIDE_TABLE_FONT_PT = 8


def write_docx(pages_blocks, output_path):
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    doc = Document()

    max_ncols = max(
        (len(row) for blocks in pages_blocks for kind, payload in blocks if kind == "table" for row in payload["grid"]),
        default=0,
    )
    is_wide = max_ncols > WIDE_TABLE_COLUMN_THRESHOLD
    if is_wide:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        for margin in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, margin, Mm(10))

    # A page whose only content was a table continuation now has an empty
    # blocks list (merge_table_continuations folded it into the previous
    # page's table) — skip it entirely rather than emitting a blank page
    # with a forced page break, which would show up as pointless empty
    # pages breaking up what's now one continuous table.
    wrote_any_page = False
    for blocks in pages_blocks:
        if not blocks:
            continue
        if wrote_any_page:
            doc.add_page_break()
        wrote_any_page = True
        for kind, payload in blocks:
            if kind == "table":
                rows = payload["grid"]
                spans = payload.get("spans", [])
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                table.autofit = True
                for r_idx, row_cells in enumerate(rows):
                    for c_idx in range(ncols):
                        text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell = table.cell(r_idx, c_idx)
                        cell.text = text
                        if is_wide:
                            # Small, consistent font across every cell so a
                            # wide table's columns stay narrow-but-readable
                            # instead of Word wrapping each cell onto many
                            # lines at its default ~11pt.
                            for para in cell.paragraphs:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                for run in para.runs:
                                    run.font.size = Pt(WIDE_TABLE_FONT_PT)
                # Reproduce the source's real merged cells (a table-header
                # spanning several columns, a label spanning several rows)
                # instead of leaving every cell separate — an unmerged table
                # visually diverges from the source even when every value is
                # individually correct, which reads as "shifted"/"messy".
                # The recognized spans can overlap each other on messy
                # real-world tables (two spans both claiming a cell) — that
                # corrupts python-docx's internal grid-offset bookkeeping if
                # applied blindly, so track which cells are already merged
                # and skip (rather than crash on) anything that overlaps.
                merged_cells = set()
                for r0, c0, r1, c1 in sorted(spans, key=lambda s: (s[0], s[1])):
                    if r1 >= len(rows) or c1 >= ncols or (r0, c0) == (r1, c1):
                        continue
                    span_cells = {(r, c) for r in range(r0, r1 + 1) for c in range(c0, c1 + 1)}
                    if span_cells & merged_cells:
                        continue
                    try:
                        table.cell(r0, c0).merge(table.cell(r1, c1))
                        merged_cells |= span_cells
                    except ValueError:
                        # A prior merge already reshaped this row/column's
                        # cell layout in a way this span no longer maps
                        # cleanly onto — skip it rather than losing the rest
                        # of the table.
                        continue
            else:
                doc.add_paragraph(payload)
    doc.save(output_path)


def main() -> int:
    if len(sys.argv) < 3:
        print("usage: ocr_layout_to_docx.py <input> <output.docx> [--pages 0,2,5]", file=sys.stderr)
        return 2

    input_path, output_path = sys.argv[1], sys.argv[2]
    pages = None
    if "--pages" in sys.argv:
        idx = sys.argv.index("--pages")
        pages = [int(p) for p in sys.argv[idx + 1].split(",") if p.strip() != ""]

    ocr = make_ocr()
    layout_detector = make_layout_detector()
    table_pipeline = LazyModel(make_table_pipeline)

    try:
        pages_blocks = ocr_layout_docx_pages(ocr, layout_detector, table_pipeline, input_path, pages)
    except ValueError as e:
        print(str(e), file=sys.stderr)
        return 2

    if not any(pages_blocks):
        print("no text detected", file=sys.stderr)
        return 1

    write_docx(pages_blocks, output_path)
    print("OK")
    return 0


if __name__ == "__main__":
    sys.exit(main())
